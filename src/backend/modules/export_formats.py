#!/usr/bin/env python3
"""
Export Formats Module
Handles exporting security reports to multiple formats: Excel, CSV, DOCX, Markdown
"""

import os
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)


class ReportExporter:
    """
    Handles exporting security reports to multiple formats.
    """

    def __init__(self, output_dir: Optional[str] = None):
        """
        Initialize the report exporter.

        Args:
            output_dir: Directory for saving exported reports
        """
        if output_dir:
            self.output_dir = Path(output_dir)
        else:
            project_root = Path(__file__).parent.parent.parent.parent
            self.output_dir = project_root / "reports"

        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"ReportExporter initialized (output: {self.output_dir})")

    def export_to_excel(self, report_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Export report data to Excel format with multiple sheets.

        Args:
            report_data: Report data dictionary
            filename: Optional filename (auto-generated if not provided)

        Returns:
            Path to the generated Excel file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"security_report_{timestamp}.xlsx"

        filepath = self.output_dir / filename

        try:
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                # Summary sheet
                summary_data = self._prepare_summary_data(report_data)
                if summary_data:
                    df_summary = pd.DataFrame([summary_data])
                    df_summary.to_excel(writer, sheet_name='Summary', index=False)

                # Findings sheet
                if report_data.get('scan_summary', {}).get('findings'):
                    findings = report_data['scan_summary']['findings']
                    df_findings = pd.DataFrame(findings)
                    df_findings.to_excel(writer, sheet_name='Security Findings', index=False)

                # Hardening results sheet
                if report_data.get('hardening_summary', {}).get('results'):
                    results = report_data['hardening_summary']['results']
                    df_results = pd.DataFrame(results)
                    df_results.to_excel(writer, sheet_name='Hardening Results', index=False)

                # Recommendations sheet
                if report_data.get('recommendations'):
                    df_recs = pd.DataFrame(report_data['recommendations'])
                    df_recs.to_excel(writer, sheet_name='Recommendations', index=False)

                # Before/After comparison
                if report_data.get('before_after'):
                    ba_data = report_data['before_after']
                    comparison = {
                        'Metric': ['Total', 'Critical', 'High', 'Medium', 'Low'],
                        'Before': [
                            ba_data.get('before', {}).get('total', 0),
                            ba_data.get('before', {}).get('critical', 0),
                            ba_data.get('before', {}).get('high', 0),
                            ba_data.get('before', {}).get('medium', 0),
                            ba_data.get('before', {}).get('low', 0)
                        ],
                        'After': [
                            ba_data.get('after', {}).get('total', 0),
                            ba_data.get('after', {}).get('critical', 0),
                            ba_data.get('after', {}).get('high', 0),
                            ba_data.get('after', {}).get('medium', 0),
                            ba_data.get('after', {}).get('low', 0)
                        ]
                    }
                    df_comparison = pd.DataFrame(comparison)
                    df_comparison['Improvement'] = df_comparison['Before'] - df_comparison['After']
                    df_comparison.to_excel(writer, sheet_name='Before-After', index=False)

            logger.info(f"Excel report generated: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"Error generating Excel report: {e}", exc_info=True)
            raise

    def export_to_csv(self, report_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Export report data to CSV format (findings only).

        Args:
            report_data: Report data dictionary
            filename: Optional filename

        Returns:
            Path to the generated CSV file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"security_findings_{timestamp}.csv"

        filepath = self.output_dir / filename

        try:
            findings = report_data.get('scan_summary', {}).get('findings', [])
            if findings:
                df = pd.DataFrame(findings)
                df.to_csv(filepath, index=False)
            else:
                # Create empty CSV with headers
                df = pd.DataFrame(columns=['title', 'severity', 'category', 'description', 'remediation'])
                df.to_csv(filepath, index=False)

            logger.info(f"CSV report generated: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"Error generating CSV report: {e}", exc_info=True)
            raise

    def export_to_docx(self, report_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Export report data to DOCX format.

        Args:
            report_data: Report data dictionary
            filename: Optional filename

        Returns:
            Path to the generated DOCX file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"security_report_{timestamp}.docx"

        filepath = self.output_dir / filename

        try:
            doc = Document()

            # Title
            title = doc.add_heading(report_data.get('title', 'Security Report'), 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Metadata
            doc.add_paragraph(f"Generated: {report_data.get('generated_date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}")
            doc.add_paragraph()

            # Executive Summary
            doc.add_heading('Executive Summary', 1)

            compliance = report_data.get('compliance', {})
            if 'current_score' in compliance:
                p = doc.add_paragraph()
                p.add_run(f"Overall Compliance Score: ").bold = True
                p.add_run(f"{compliance['current_score']}%")

            scan_summary = report_data.get('scan_summary', {})
            if scan_summary:
                p = doc.add_paragraph()
                p.add_run(f"Total Findings: ").bold = True
                p.add_run(f"{scan_summary.get('total_findings', 0)}")

                p = doc.add_paragraph()
                p.add_run("Severity Breakdown:\n").bold = True
                p.add_run(f"• Critical: {scan_summary.get('critical', 0)}\n")
                p.add_run(f"• High: {scan_summary.get('high', 0)}\n")
                p.add_run(f"• Medium: {scan_summary.get('medium', 0)}\n")
                p.add_run(f"• Low: {scan_summary.get('low', 0)}")

            doc.add_page_break()

            # Security Findings
            if scan_summary.get('findings'):
                doc.add_heading('Security Findings', 1)

                for finding in scan_summary['findings']:
                    doc.add_heading(finding.get('title', 'Unknown'), 2)

                    p = doc.add_paragraph()
                    p.add_run('Severity: ').bold = True
                    severity_run = p.add_run(finding.get('severity', 'Unknown'))
                    # Color code severity
                    if finding.get('severity', '').lower() == 'critical':
                        severity_run.font.color.rgb = RGBColor(220, 53, 69)
                    elif finding.get('severity', '').lower() == 'high':
                        severity_run.font.color.rgb = RGBColor(255, 107, 107)

                    p = doc.add_paragraph()
                    p.add_run('Category: ').bold = True
                    p.add_run(finding.get('category', 'General'))

                    p = doc.add_paragraph()
                    p.add_run('Description: ').bold = True
                    p.add_run(finding.get('description', 'No description available'))

                    if finding.get('remediation'):
                        p = doc.add_paragraph()
                        p.add_run('Remediation: ').bold = True
                        p.add_run(finding['remediation'])

                    doc.add_paragraph()  # Spacing

                doc.add_page_break()

            # Hardening Results
            hardening_summary = report_data.get('hardening_summary', {})
            if hardening_summary.get('results'):
                doc.add_heading('Remediation Actions Taken', 1)

                p = doc.add_paragraph()
                p.add_run(f"Total Rules: ").bold = True
                p.add_run(f"{hardening_summary.get('total_rules', 0)}\n")
                p.add_run(f"Successful: ").bold = True
                p.add_run(f"{hardening_summary.get('successful_rules', 0)}\n")
                p.add_run(f"Failed: ").bold = True
                p.add_run(f"{hardening_summary.get('failed_rules', 0)}")

                doc.add_paragraph()

                for result in hardening_summary['results']:
                    doc.add_heading(result.get('rule_name', 'Unknown Rule'), 3)

                    p = doc.add_paragraph()
                    p.add_run('Status: ').bold = True
                    status_run = p.add_run(result.get('status', 'Unknown'))

                    if result.get('before_value') and result.get('after_value'):
                        p = doc.add_paragraph()
                        p.add_run('Change: ').bold = True
                        p.add_run(f"{result['before_value']} → {result['after_value']}")

                    if result.get('error'):
                        p = doc.add_paragraph()
                        p.add_run('Error: ').bold = True
                        p.add_run(result['error'])

                    doc.add_paragraph()  # Spacing

                doc.add_page_break()

            # Recommendations
            if report_data.get('recommendations'):
                doc.add_heading('Recommendations', 1)

                for rec in report_data['recommendations']:
                    doc.add_heading(rec.get('title', 'Recommendation'), 2)
                    p = doc.add_paragraph()
                    p.add_run('Priority: ').bold = True
                    p.add_run(rec.get('priority', 'medium').upper())

                    p = doc.add_paragraph(rec.get('description', ''))
                    doc.add_paragraph()

            # Footer
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"System Hardening Tool - Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.save(filepath)
            logger.info(f"DOCX report generated: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"Error generating DOCX report: {e}", exc_info=True)
            raise

    def export_to_markdown(self, report_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Export report data to Markdown format.

        Args:
            report_data: Report data dictionary
            filename: Optional filename

        Returns:
            Path to the generated Markdown file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"security_report_{timestamp}.md"

        filepath = self.output_dir / filename

        try:
            md_content = []

            # Title
            md_content.append(f"# {report_data.get('title', 'Security Report')}\n")
            md_content.append(f"**Generated:** {report_data.get('generated_date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}\n")
            md_content.append("---\n")

            # Executive Summary
            md_content.append("## Executive Summary\n")

            compliance = report_data.get('compliance', {})
            if 'current_score' in compliance:
                md_content.append(f"**Overall Compliance Score:** {compliance['current_score']}%\n")

            scan_summary = report_data.get('scan_summary', {})
            if scan_summary:
                md_content.append(f"\n**Total Findings:** {scan_summary.get('total_findings', 0)}\n")
                md_content.append("\n### Severity Breakdown\n")
                md_content.append(f"- 🔴 **Critical:** {scan_summary.get('critical', 0)}\n")
                md_content.append(f"- 🟠 **High:** {scan_summary.get('high', 0)}\n")
                md_content.append(f"- 🟡 **Medium:** {scan_summary.get('medium', 0)}\n")
                md_content.append(f"- 🔵 **Low:** {scan_summary.get('low', 0)}\n")

            # Security Findings
            if scan_summary.get('findings'):
                md_content.append("\n## Security Findings\n")

                for idx, finding in enumerate(scan_summary['findings'], 1):
                    severity_icon = {
                        'critical': '🔴',
                        'high': '🟠',
                        'medium': '🟡',
                        'low': '🔵',
                        'info': '⚪'
                    }.get(finding.get('severity', 'info').lower(), '⚪')

                    md_content.append(f"\n### {idx}. {finding.get('title', 'Unknown')} {severity_icon}\n")
                    md_content.append(f"**Severity:** {finding.get('severity', 'Unknown')}\n")
                    md_content.append(f"**Category:** {finding.get('category', 'General')}\n")
                    md_content.append(f"\n**Description:**\n{finding.get('description', 'No description available')}\n")

                    if finding.get('remediation'):
                        md_content.append(f"\n**Remediation:**\n{finding['remediation']}\n")

            # Hardening Results
            hardening_summary = report_data.get('hardening_summary', {})
            if hardening_summary.get('results'):
                md_content.append("\n## Remediation Actions Taken\n")
                md_content.append(f"- **Total Rules:** {hardening_summary.get('total_rules', 0)}\n")
                md_content.append(f"- **Successful:** {hardening_summary.get('successful_rules', 0)} ✅\n")
                md_content.append(f"- **Failed:** {hardening_summary.get('failed_rules', 0)} ❌\n")

                md_content.append("\n### Detailed Results\n")

                for result in hardening_summary['results']:
                    status_icon = {'success': '✅', 'failed': '❌', 'skipped': '⊘'}.get(result.get('status', ''), '○')

                    md_content.append(f"\n#### {result.get('rule_name', 'Unknown Rule')} {status_icon}\n")
                    md_content.append(f"**Status:** {result.get('status', 'Unknown')}\n")

                    if result.get('before_value') and result.get('after_value'):
                        md_content.append(f"**Change:** `{result['before_value']}` → `{result['after_value']}`\n")

                    if result.get('error'):
                        md_content.append(f"**Error:** {result['error']}\n")

            # Before/After Comparison
            if report_data.get('before_after'):
                ba_data = report_data['before_after']
                md_content.append("\n## Before/After Comparison\n")
                md_content.append("\n| Severity | Before | After | Improvement |\n")
                md_content.append("|----------|--------|-------|-------------|\n")

                metrics = ['total', 'critical', 'high', 'medium', 'low']
                for metric in metrics:
                    before_val = ba_data.get('before', {}).get(metric, 0)
                    after_val = ba_data.get('after', {}).get(metric, 0)
                    improvement = before_val - after_val
                    md_content.append(f"| {metric.capitalize()} | {before_val} | {after_val} | {improvement} |\n")

            # Recommendations
            if report_data.get('recommendations'):
                md_content.append("\n## Recommendations\n")

                for rec in report_data['recommendations']:
                    priority_icon = {
                        'critical': '🔴',
                        'high': '🟠',
                        'medium': '🟡',
                        'low': '🔵',
                        'info': '💡'
                    }.get(rec.get('priority', 'info'), '💡')

                    md_content.append(f"\n### {priority_icon} {rec.get('title', 'Recommendation')}\n")
                    md_content.append(f"{rec.get('description', '')}\n")

            # Footer
            md_content.append("\n---\n")
            md_content.append(f"*Report generated by System Hardening Tool on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n")

            # Write to file
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(''.join(md_content))

            logger.info(f"Markdown report generated: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"Error generating Markdown report: {e}", exc_info=True)
            raise

    def _prepare_summary_data(self, report_data: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare summary data for Excel export."""
        summary = {
            'Report Title': report_data.get('title', 'Security Report'),
            'Generated Date': report_data.get('generated_date', ''),
            'Compliance Score': report_data.get('compliance', {}).get('current_score', 0)
        }

        scan_summary = report_data.get('scan_summary', {})
        if scan_summary:
            summary['Total Findings'] = scan_summary.get('total_findings', 0)
            summary['Critical'] = scan_summary.get('critical', 0)
            summary['High'] = scan_summary.get('high', 0)
            summary['Medium'] = scan_summary.get('medium', 0)
            summary['Low'] = scan_summary.get('low', 0)

        hardening_summary = report_data.get('hardening_summary', {})
        if hardening_summary:
            summary['Total Rules Applied'] = hardening_summary.get('total_rules', 0)
            summary['Successful Rules'] = hardening_summary.get('successful_rules', 0)
            summary['Failed Rules'] = hardening_summary.get('failed_rules', 0)

        return summary


# Example usage and testing
if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

    print("=" * 70)
    print("Report Exporter - Test Run")
    print("=" * 70)

    # Sample data
    sample_data = {
        "title": "Test Security Report",
        "generated_date": "2025-11-23 17:00:00",
        "compliance": {"current_score": 85},
        "scan_summary": {
            "total_findings": 3,
            "critical": 1,
            "high": 1,
            "medium": 1,
            "low": 0,
            "findings": [
                {
                    "title": "SSH Root Login Enabled",
                    "severity": "critical",
                    "category": "Network Security",
                    "description": "Root login is enabled",
                    "remediation": "Disable root login"
                }
            ]
        }
    }

    exporter = ReportExporter()

    print("\nTesting Excel export...")
    excel_path = exporter.export_to_excel(sample_data)
    print(f"✓ Excel exported: {excel_path}")

    print("\nTesting CSV export...")
    csv_path = exporter.export_to_csv(sample_data)
    print(f"✓ CSV exported: {csv_path}")

    print("\nTesting DOCX export...")
    docx_path = exporter.export_to_docx(sample_data)
    print(f"✓ DOCX exported: {docx_path}")

    print("\nTesting Markdown export...")
    md_path = exporter.export_to_markdown(sample_data)
    print(f"✓ Markdown exported: {md_path}")

    print("\n" + "=" * 70)
    print("All exports completed successfully!")
    print("=" * 70)
